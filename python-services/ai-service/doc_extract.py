from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
from typing import Optional, List
import io
import csv

# DOCX support (python-docx)
try:
    import docx  # type: ignore
    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover
    DOCX_AVAILABLE = False

router = APIRouter()

class TextExtractResponse(BaseModel):
    text: Optional[str] = None


def _normalize(text: str, max_len: int = 20000) -> str:
    t = " ".join((text or "").split())
    return t[:max_len]


@router.post("/extract/docx", response_model=TextExtractResponse)
async def extract_docx_endpoint(file: UploadFile = File(...)):
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=503, detail="DOCX extraction library not available")
    if file.content_type not in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    ) and not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail=f"Unsupported content type: {file.content_type}")

    try:
        data = await file.read()
        document = docx.Document(io.BytesIO(data))
        parts: List[str] = []
        # Paragraphs
        for p in document.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        # Tables
        for tbl in document.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return TextExtractResponse(text=_normalize(text))
    except Exception as e:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Failed to extract DOCX text: {str(e)}")


@router.post("/extract/csv", response_model=TextExtractResponse)
async def extract_csv_endpoint(file: UploadFile = File(...)):
    if file.content_type not in (
        "text/csv",
        "application/csv",
        "application/vnd.ms-excel",
        "application/octet-stream",
    ) and not (file.filename or "").lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail=f"Unsupported content type: {file.content_type}")

    try:
        data = await file.read()
        # Decode with fallback
        text = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = []
        max_rows = 200  # cap rows to keep prompt small
        for i, row in enumerate(reader):
            if i == 0:
                rows.append("Header: " + ", ".join(row))
            else:
                rows.append("Row " + str(i) + ": " + ", ".join(row))
            if i >= max_rows:
                break
        return TextExtractResponse(text=_normalize("\n".join(rows)))
    except Exception as e:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Failed to extract CSV text: {str(e)}")

